#!/usr/bin/env python3
"""
Script pour créer un document Word formaté à partir du guide MOBIQ
Nécessite: pip install python-docx
"""

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

def create_mobiq_word_doc():
    # Créer un nouveau document
    doc = Document()
    
    # Titre principal
    title = doc.add_heading('MOBIQ - Executable Installation and User Guide', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Sous-titre
    subtitle = doc.add_paragraph('MOBIQ Desktop Application')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].bold = True
    
    # Description
    desc = doc.add_paragraph('Complete guide for MOBIQ executable installation and usage')
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    desc.runs[0].italic = True
    
    # Features
    features = doc.add_paragraph('Simple Installation • Electron Interface • Ready to Use')
    features.alignment = WD_ALIGN_PARAGRAPH.CENTER
    features.runs[0].bold = True
    
    doc.add_page_break()
    
    # Table des matières
    doc.add_heading('Table of Contents', level=1)
    
    toc_items = [
        "Download and Installation",
        "First Launch", 
        "Initial Configuration",
        "Device Connection",
        "Electron Interface",
        "Using Tests",
        "Automated Workflows",
        "Monitoring and Reports",
        "Settings and Preferences",
        "Troubleshooting"
    ]
    
    for i, item in enumerate(toc_items, 1):
        toc_para = doc.add_paragraph(f"{i}. {item}")
        toc_para.style = 'List Number'
    
    doc.add_page_break()
    
    # Section 1: Download and Installation
    doc.add_heading('Download and Installation', level=1)
    
    doc.add_heading('System Requirements', level=2)
    
    doc.add_heading('Minimum:', level=3)
    min_req = [
        "OS: Windows 10 (64-bit) or higher",
        "RAM: 4 GB minimum (8 GB recommended)",
        "Storage: 500 MB free space",
        "USB: USB 2.0 port or higher for Android connection"
    ]
    
    for req in min_req:
        p = doc.add_paragraph(req)
        p.style = 'List Bullet'
    
    doc.add_heading('Recommended:', level=3)
    rec_req = [
        "OS: Windows 11",
        "RAM: 8 GB or more", 
        "Storage: 1 GB free space",
        "USB: USB 3.0 port for better performance"
    ]
    
    for req in rec_req:
        p = doc.add_paragraph(req)
        p.style = 'List Bullet'
    
    # Executable Download
    doc.add_heading('1. Executable Download', level=2)
    
    doc.add_paragraph('[📸 Insert download page screenshot here]').runs[0].italic = True
    doc.add_paragraph('Suggested photo location: media/screenshots/download-page.png').runs[0].italic = True
    
    download_steps = [
        "Access the download page:\n  • Official website: https://mobiq-framework.com/download\n  • Or from GitHub Releases",
        "Select version:\n  • MOBIQ-Setup-v2.2.0.exe (Full version with installer)\n  • MOBIQ-Portable-v2.2.0.zip (Portable version)",
        "Verify integrity:\n  • SHA256 checksum provided on page\n  • Verified digital signature"
    ]
    
    for i, step in enumerate(download_steps, 1):
        p = doc.add_paragraph(f"{i}. {step}")
    
    # Installation with Installer
    doc.add_heading('2. Installation with Installer', level=2)
    
    doc.add_paragraph('[📸 Insert installer screenshot here]').runs[0].italic = True
    doc.add_paragraph('Suggested photo location: media/screenshots/installer-welcome.png').runs[0].italic = True
    
    # Installation steps
    install_steps = [
        ("Step 1: Installer Launch", [
            "Run MOBIQ-Setup-v2.2.0.exe as administrator",
            "Accept Windows security warning if necessary", 
            "Click 'Next' on the welcome screen"
        ]),
        ("Step 2: License and Terms", [
            "Read the MIT license terms",
            "Check 'I accept the license terms'",
            "Click 'Next'"
        ]),
        ("Step 3: Installation Directory Choice", [
            "Default directory: C:\\Program Files\\MOBIQ Framework\\",
            "Customize if needed with 'Browse'",
            "Check available space (minimum 500 MB)",
            "Click 'Next'"
        ]),
        ("Step 4: Components to Install", [
            "✓ Main application (required)",
            "✓ ADB drivers (recommended)",
            "✓ Test modules (required)",
            "☐ Examples and tutorials (optional)",
            "☐ Offline documentation (optional)"
        ]),
        ("Step 5: Shortcuts and Options", [
            "✓ Create Desktop shortcut",
            "✓ Add to Start Menu",
            "✓ Create Taskbar shortcut", 
            "☐ Launch MOBIQ at Windows startup",
            "✓ Associate .mobiq files with application"
        ]),
        ("Step 6: Installation", [
            "Click 'Install'",
            "Wait for installation completion (2-5 minutes)",
            "Progress displayed in real-time:\n  • File extraction\n  • ADB drivers installation\n  • Service configuration\n  • Shortcut creation"
        ]),
        ("Step 7: Completion", [
            "Installation completed successfully",
            "Final options:\n  ✓ Launch MOBIQ now\n  ☐ Show release notes",
            "Click 'Finish'"
        ])
    ]
    
    for step_title, step_items in install_steps:
        doc.add_heading(step_title, level=3)
        
        # Ajouter screenshot placeholder
        doc.add_paragraph(f'[📸 Insert {step_title.lower()} screenshot here]').runs[0].italic = True
        doc.add_paragraph(f'Suggested photo location: media/screenshots/installer-{step_title.split(":")[0].split()[-1].lower()}.png').runs[0].italic = True
        
        for i, item in enumerate(step_items, 1):
            p = doc.add_paragraph(f"{i}. {item}")
    
    # Portable Installation
    doc.add_heading('3. Portable Installation (Alternative)', level=2)
    
    portable_steps = [
        "Extract MOBIQ-Portable-v2.2.0.zip",
        "Run MOBIQ.exe directly", 
        "No system installation required"
    ]
    
    doc.add_paragraph("For portable version:")
    for i, step in enumerate(portable_steps, 1):
        p = doc.add_paragraph(f"{i}. {step}")
    
    doc.add_page_break()
    
    # Section 2: First Launch
    doc.add_heading('First Launch', level=1)
    
    doc.add_heading('1. Startup Screen', level=2)
    doc.add_paragraph('[📸 Insert startup screen screenshot here]').runs[0].italic = True
    doc.add_paragraph('Suggested photo location: media/screenshots/startup-screen.png').runs[0].italic = True
    
    doc.add_paragraph("On first launch, MOBIQ displays:")
    startup_items = [
        "Animated logo with progress bar",
        "Component verification:\n  ✓ ADB Engine\n  ✓ Test Modules\n  ✓ User Interface\n  ✓ Local Database"
    ]
    
    for item in startup_items:
        p = doc.add_paragraph(item)
        p.style = 'List Bullet'
    
    # Configuration Wizard
    doc.add_heading('2. Initial Configuration Wizard', level=2)
    doc.add_paragraph('[📸 Insert wizard screenshot here]').runs[0].italic = True
    doc.add_paragraph('Suggested photo location: media/screenshots/setup-wizard.png').runs[0].italic = True
    
    wizard_steps = [
        ("Step 1: Welcome", [
            "Welcome message to MOBIQ",
            "Overview of main features",
            "Click 'Start Configuration'"
        ]),
        ("Step 2: ADB Verification", [
            "Automatic test for ADB presence",
            "Possible results:\n  ✓ ADB detected: Version X.X.X found\n  ✗ ADB missing: Automatic installation offered\n  ⚠ Outdated version: Update recommended",
            "Automatic actions:\n  • ADB installation if missing\n  • System PATH configuration\n  • Functionality test"
        ]),
        ("Step 3: User Preferences", [
            "Language: French / English",
            "Theme: Light / Dark / Automatic",
            "Notifications: Enabled / Disabled",
            "Auto-start: Yes / No",
            "Working folder: Directory choice for reports"
        ]),
        ("Step 4: Connectivity Test", [
            "Message: 'Connect an Android device to test'",
            "Instructions:\n  • Enable USB debugging\n  • Connect via USB cable\n  • Authorize on device",
            "Automatic test as soon as device is detected"
        ])
    ]
    
    for step_title, step_items in wizard_steps:
        doc.add_heading(step_title, level=3)
        
        # Screenshot placeholder
        step_name = step_title.split(":")[1].strip().lower().replace(" ", "-")
        doc.add_paragraph(f'[📸 Insert {step_name} screenshot here]').runs[0].italic = True
        doc.add_paragraph(f'Suggested photo location: media/screenshots/{step_name}.png').runs[0].italic = True
        
        for i, item in enumerate(step_items, 1):
            if step_title == "Step 3: User Preferences":
                p = doc.add_paragraph(f"• {item}")
            else:
                p = doc.add_paragraph(f"{i}. {item}")
    
    # Ajouter note de fin
    doc.add_page_break()
    doc.add_heading('Support and Resources', level=1)
    
    support_items = [
        "Complete Documentation: Accessible via F1 in application",
        "Video Tutorials: Integrated in help",
        "Community Forum: https://community.mobiq-framework.com", 
        "Technical Support: support@mobiq-framework.com"
    ]
    
    doc.add_paragraph("Available Resources:")
    for i, item in enumerate(support_items, 1):
        p = doc.add_paragraph(f"{i}. {item}")
    
    # Release Notes
    doc.add_heading('Release Notes', level=1)
    
    doc.add_paragraph("MOBIQ Desktop v2.2.0").runs[0].bold = True
    
    features = [
        "Modernized Electron interface",
        "29 integrated test modules", 
        "Visual drag-and-drop workflows",
        "Automatic PDF/Excel reports",
        "Enhanced multi-device support"
    ]
    
    for feature in features:
        p = doc.add_paragraph(feature)
        p.style = 'List Bullet'
    
    doc.add_paragraph("Compatibility:").runs[0].bold = True
    compatibility = [
        "Windows 10/11 (64-bit)",
        "Android 7.0+ (API 24+)",
        "ADB version 1.0.39+"
    ]
    
    for comp in compatibility:
        p = doc.add_paragraph(comp)
        p.style = 'List Bullet'
    
    # Note finale
    doc.add_paragraph("\nThis guide covers the usage of MOBIQ Desktop executable. For development or source installation, consult the separate developer guide.").runs[0].italic = True
    
    # Sauvegarder le document
    doc.save('MOBIQ_User_Guide_Complete.docx')
    print("✅ Document Word créé avec succès : MOBIQ_User_Guide_Complete.docx")

if __name__ == "__main__":
    try:
        create_mobiq_word_doc()
    except ImportError:
        print("❌ Erreur : python-docx n'est pas installé")
        print("📦 Installez avec : pip install python-docx")
    except Exception as e:
        print(f"❌ Erreur lors de la création du document : {e}")